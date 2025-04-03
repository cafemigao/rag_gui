import os
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from threading import Thread
from queue import Queue
import torch
import fitz
from docx import Document as DocxDocument
import openpyxl
import requests
from urllib.parse import urlparse
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
from langchain_huggingface import HuggingFacePipeline
from langchain.schema import Document
import gc
from multiprocessing import cpu_count
from gtts import gTTS

DEVICES = "cuda" if torch.cuda.is_available() else "cpu"
MODEL_NAME = "mistralai/Mistral-7B-Instruct-v0.2"
EMBEDDING_MODEL = "sentence-transformers/paraphrase-MiniLM-L3-v2"
CPU_CORES = cpu_count()

class RAGGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("RAG GUI *cafemigao*")
        self.root.geometry("900x600")
        self.root.configure(bg="#F0F4F8")  # Fondo claro
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.vector_store = None
        self.llm = None
        self.log_queue = Queue()
        self.is_processing = False

        # Estilo
        style = ttk.Style()
        style.configure("TButton", font=("Arial", 10, "bold"), padding=5)
        style.configure("TLabel", font=("Arial", 11), background="#F0F4F8")
        style.map("TButton", background=[("active", "#A3BFFA"), ("disabled", "#D3D3D3")])

        # Frame principal
        main_frame = ttk.Frame(root, padding=10, style="Main.TFrame")
        main_frame.grid(row=0, column=0, sticky="nsew")
        style.configure("Main.TFrame", background="#F0F4F8")

        # Área de logs
        log_frame = ttk.LabelFrame(main_frame, text="Registro", padding=5)
        log_frame.grid(row=0, column=0, columnspan=4, padx=5, pady=5, sticky="nsew")
        self.log_area = scrolledtext.ScrolledText(
            log_frame, wrap=tk.WORD, width=100, height=10, bg="#2D2D2D", fg="#FFFFFF",
            font=("Consolas", 10), insertbackground="white"
        )
        self.log_area.pack(fill="both", expand=True)
        self.log_handler = TextHandler(self.log_area, self.log_queue)
        self.root.after(100, self.update_logs)
        self.log_message(f"🔹 Dispositivo: {DEVICES.upper()} | Núcleos: {CPU_CORES}")

        # Frame de entrada de fuente
        source_frame = ttk.LabelFrame(main_frame, text="Cargar Documento", padding=5)
        source_frame.grid(row=1, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
        ttk.Label(source_frame, text="Ruta o URL:").pack(side="left", padx=5)
        self.source_entry = ttk.Entry(source_frame, width=50, font=("Arial", 10))
        self.source_entry.pack(side="left", fill="x", expand=True, padx=5)
        self.browse_button = ttk.Button(source_frame, text="📂 Seleccionar", command=self.select_file)
        self.browse_button.pack(side="left", padx=5)
        self.process_button = ttk.Button(source_frame, text="▶ Procesar", command=self.process_source, style="Accent.TButton")
        self.process_button.pack(side="left", padx=5)
        style.configure("Accent.TButton", background="#4C78A8")

        # Frame de pregunta
        query_frame = ttk.LabelFrame(main_frame, text="Consulta", padding=5)
        query_frame.grid(row=2, column=0, columnspan=4, padx=5, pady=5, sticky="ew")
        ttk.Label(query_frame, text="Pregunta:").pack(side="left", padx=5)
        self.query_entry = ttk.Entry(query_frame, width=50, font=("Arial", 10))
        self.query_entry.pack(side="left", fill="x", expand=True, padx=5)
        self.ask_button = ttk.Button(query_frame, text="❓ Enviar", command=self.ask_question, state="disabled")
        self.ask_button.pack(side="left", padx=5)

        # Frame de respuesta
        answer_frame = ttk.LabelFrame(main_frame, text="Respuesta", padding=5)
        answer_frame.grid(row=3, column=0, columnspan=4, padx=5, pady=5, sticky="nsew")
        self.answer_area = scrolledtext.ScrolledText(
            answer_frame, wrap=tk.WORD, width=100, height=10, bg="#FFFFFF", fg="#333333",
            font=("Arial", 10), insertbackground="black"
        )
        self.answer_area.pack(fill="both", expand=True)

        # Frame de controles
        control_frame = ttk.Frame(main_frame, padding=5)
        control_frame.grid(row=4, column=0, columnspan=4, pady=5, sticky="ew")
        self.speak_button = ttk.Button(control_frame, text="🔊 Leer", command=self.speak_answer, state="disabled")
        self.speak_button.pack(side="left", padx=5)
        self.copy_button = ttk.Button(control_frame, text="📋 Copiar", command=self.copy_answer, state="disabled")
        self.copy_button.pack(side="left", padx=5)
        self.exit_button = ttk.Button(control_frame, text="✖ Salir", command=self.on_closing, style="Exit.TButton")
        self.exit_button.pack(side="right", padx=5)
        style.configure("Exit.TButton", background="#E57373")

        # Configurar expansión
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(0, weight=1)
        main_frame.grid_rowconfigure(3, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)

    def log_message(self, message):
        self.log_queue.put(message)

    def update_logs(self):
        while not self.log_queue.empty():
            msg = self.log_queue.get()
            self.log_area.insert(tk.END, msg + "\n")
            self.log_area.see(tk.END)
        self.root.after(100, self.update_logs)

    def on_closing(self):
        if messagebox.askokcancel("Salir", "¿Seguro que quieres salir?"):
            if self.llm or self.vector_store:
                del self.llm, self.vector_store
                gc.collect()
            self.root.destroy()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[
            ("Todos los archivos", "*.*"), ("Texto", "*.txt"), ("PDF", "*.pdf"),
            ("Word", "*.docx"), ("Excel", "*.xlsx", "*.log")
        ])
        if file_path:
            self.source_entry.delete(0, tk.END)
            self.source_entry.insert(0, file_path)

    def process_source(self):
        if self.is_processing:
            return
        self.is_processing = True
        source = self.source_entry.get().strip()
        if not source:
            messagebox.showwarning("Advertencia", "Por favor, ingrese una URL o seleccione un archivo.")
            self.is_processing = False
            return

        self.process_button.config(text="Procesando...", state="disabled")
        self.log_message("⏳ Procesando documento...")

        def process():
            try:
                documents = self.load_documents(source)
                if not documents:
                    self.log_message("❌ No se pudo procesar el documento.")
                    self.reset_process_button()
                    return

                texts = self.split_documents(documents)
                self.log_message(f"📝 Fragmentos generados: {len(texts)}")

                embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL, model_kwargs={"device": DEVICES})
                self.vector_store = FAISS.from_documents(texts, embeddings)
                del texts, embeddings
                gc.collect()

                self.setup_llm()
                self.root.after(0, lambda: self.ask_button.config(state="normal"))
                self.root.after(0, lambda: self.reset_process_button())
                self.log_message("✅ Listo para preguntas.")
            except Exception as e:
                self.log_message(f"❌ Error al procesar: {e}")
                self.root.after(0, lambda: self.reset_process_button())

        Thread(target=process, daemon=True).start()

    def reset_process_button(self):
        self.process_button.config(text="▶ Procesar", state="normal")
        self.is_processing = False

    def load_documents(self, source):
        documents = []
        try:
            if source.startswith("http"):
                response = requests.get(source, timeout=10)
                response.raise_for_status()
                text = response.text
                documents.append(Document(page_content=text, metadata={"source": source}))
            elif os.path.exists(source):
                if source.endswith('.pdf'):
                    with fitz.open(source) as doc:
                        text = "\n".join([page.get_text("text") for page in doc])
                    documents.append(Document(page_content=text, metadata={"source": source}))
                elif source.endswith('.docx'):
                    doc = DocxDocument(source)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    documents.append(Document(page_content=text, metadata={"source": source}))
                elif source.endswith('.xlsx'):
                    workbook = openpyxl.load_workbook(source)
                    sheet = workbook.active
                    text = "\n".join([str(cell.value) for row in sheet.rows for cell in row if cell.value])
                    documents.append(Document(page_content=text, metadata={"source": source}))
                elif source.endswith(('.txt', '.log')):
                    loader = TextLoader(source)
                    documents.extend(loader.load())
            return documents
        except Exception as e:
            self.log_message(f"❌ Error al cargar {source}: {e}")
            return []

    def split_documents(self, documents):
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        return splitter.split_documents(documents)

    def setup_llm(self):
        tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
        if tokenizer.pad_token_id is None:
            tokenizer.pad_token_id = tokenizer.eos_token_id
        model = AutoModelForCausalLM.from_pretrained(
            MODEL_NAME, device_map=DEVICES, torch_dtype=torch.float16, low_cpu_mem_usage=True
        )
        generator = pipeline(
            "text-generation", model=model, tokenizer=tokenizer, max_new_tokens=500,
            do_sample=False, pad_token_id=tokenizer.eos_token_id
        )
        self.llm = HuggingFacePipeline(pipeline=generator)
        self.log_message("✅ Modelo de lenguaje cargado.")

    def ask_question(self):
        if self.is_processing:
            return
        self.is_processing = True
        question = self.query_entry.get().strip()
        if not question:
            messagebox.showwarning("Advertencia", "Por favor, ingrese una pregunta.")
            self.is_processing = False
            return

        self.ask_button.config(state="disabled")
        self.log_message(f"❓ Pregunta: {question}")

        def process():
            try:
                retriever = self.vector_store.as_retriever(search_kwargs={"k": 10})
                retrieved_docs = retriever.invoke(question)
                context = "\n\n".join([doc.page_content for doc in retrieved_docs])
                prompt = f"""
                [INST] Basado exclusivamente en el contenido siguiente, responde detalladamente en español:
                {context}
                Pregunta: {question}
                Respuesta:
                [/INST]
                """
                response = self.llm.invoke(prompt)
                answer = response.split("[/INST]")[-1].strip() if "[/INST]" in response else response.strip()
                self.root.after(0, lambda: self.display_answer(answer))
                self.root.after(0, lambda: self.ask_button.config(state="normal"))
                self.root.after(0, lambda: self.speak_button.config(state="normal"))
                self.root.after(0, lambda: self.copy_button.config(state="normal"))
                self.is_processing = False
            except Exception as e:
                self.log_message(f"❌ Error al responder: {e}")
                self.root.after(0, lambda: self.ask_button.config(state="normal"))
                self.is_processing = False

        Thread(target=process, daemon=True).start()

    def display_answer(self, answer):
        self.answer_area.delete("1.0", tk.END)
        self.answer_area.insert(tk.END, answer)

    def speak_answer(self):
        answer_text = self.answer_area.get("1.0", tk.END).strip()
        if not answer_text:
            messagebox.showwarning("Advertencia", "No hay texto para leer.")
            return

        self.speak_button.config(state="disabled")
        self.log_message("🔊 Generando audio...")

        def speak():
            audio_file = "answer_audio.mp3"
            try:
                tts = gTTS(text=answer_text, lang="es", slow=False)
                tts.save(audio_file)
                os.system(f"mpg123 -q {audio_file}")
                self.root.after(0, lambda: self.log_message("✅ Audio reproducido."))
            except Exception as e:
                self.root.after(0, lambda: self.log_message(f"❌ Error de audio: {e}. Verifica mpg123 y el sonido."))
            finally:
                if os.path.exists(audio_file):
                    os.remove(audio_file)
                self.root.after(0, lambda: self.speak_button.config(state="normal"))

        Thread(target=speak, daemon=True).start()

    def copy_answer(self):
        answer_text = self.answer_area.get("1.0", tk.END).strip()
        if not answer_text:
            messagebox.showwarning("Advertencia", "No hay texto para copiar.")
            return
        self.root.clipboard_clear()
        self.root.clipboard_append(answer_text)
        self.log_message("📋 Respuesta copiada.")

class TextHandler:
    def __init__(self, text_widget, queue):
        self.text_widget = text_widget
        self.queue = queue

    def emit(self, record):
        self.queue.put(record)

if __name__ == "__main__":
    root = tk.Tk()
    app = RAGGUI(root)
    root.mainloop()
